from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook


def extract_pdf_text_by_page(filepath: str) -> list[tuple[int, str]]:
    """
    Returns a list of (page_number, text) tuples for a PDF file.
    """
    reader = PdfReader(filepath)
    pages = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append((page_num, text))
    return pages


def extract_docx_text_by_page(filepath: str) -> list[tuple[int, str]]:
    """
    Extracts both paragraphs AND tables from a Word document.

    IMPORTANT: python-docx doc.paragraphs does NOT include table cells —
    tables are a separate doc.tables collection. Without this fix, any
    register/matrix-style .docx (like Project_Risk_Register.docx) would
    return only the heading text, dropping all actual table data silently.
    """
    doc = DocxDocument(filepath)

    # Extract non-empty paragraphs
    paragraph_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    # Extract tables — critical for register/matrix documents
    # Each row becomes a pipe-separated line for readability in chunks
    table_parts = []
    for table_index, table in enumerate(doc.tables, start=1):
        table_parts.append(f"\n[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # skip completely empty rows
                table_parts.append(" | ".join(cells))

    full_text = paragraph_text
    if table_parts:
        full_text += "\n\n" + "\n".join(table_parts)

    return [(1, full_text)]


def extract_xlsx_text_by_page(filepath: str) -> list[tuple[int, str]]:
    """
    Each sheet in the workbook is treated as one 'page'.
    Each row is converted into a readable comma-separated line of text.
    """
    workbook = load_workbook(filepath, data_only=True)
    pages = []

    for sheet_index, sheet_name in enumerate(workbook.sheetnames, start=1):
        sheet = workbook[sheet_name]
        lines = [f"Sheet: {sheet_name}"]

        for row in sheet.iter_rows(values_only=True):
            # Skip completely empty rows
            if all(cell is None for cell in row):
                continue
            row_text = ", ".join(str(cell) if cell is not None else "" for cell in row)
            lines.append(row_text)

        sheet_text = "\n".join(lines)
        pages.append((sheet_index, sheet_text))

    return pages


def extract_text_by_page(filepath: str, filename: str) -> list[tuple[int, str]]:
    """
    Detects file type by extension and routes to the correct extractor.
    """
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        return extract_pdf_text_by_page(filepath)
    elif lower_name.endswith(".docx"):
        return extract_docx_text_by_page(filepath)
    elif lower_name.endswith(".xlsx"):
        return extract_xlsx_text_by_page(filepath)
    else:
        raise ValueError(f"Unsupported file type: {filename}")